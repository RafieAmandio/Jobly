import io
import logging
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

TEMPLATE_DIR = Path(__file__).parent.parent / "templates"

# Palette/format mirrors the owner's master CV (serif body, centred header,
# uppercase navy section titles with rules, company-left / dates-right rows).
NAVY = RGBColor(0x1F, 0x4E, 0x79)
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)
RIGHT_TAB_INCHES = 7.0


def _contact_line(contact: dict | None) -> list[str]:
    """Ordered, non-empty contact fields for the centred header line."""
    if not contact:
        return []
    parts = []
    for key in ("location", "email", "phone", "linkedin"):
        value = contact.get(key)
        if value:
            parts.append(str(value).strip())
    return parts


# --------------------------------------------------------------------------- #
# DOCX helpers
# --------------------------------------------------------------------------- #
def _set_bottom_border(paragraph, size: int = 6, color: str = "BFBFBF") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def _section_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = NAVY
    _set_bottom_border(p)


def _entry_header(doc, org: str, period: str) -> None:
    """Org bold on the left, period right-aligned on the same line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(RIGHT_TAB_INCHES), WD_TAB_ALIGNMENT.RIGHT
    )
    org_run = p.add_run(org)
    org_run.bold = True
    if period:
        p.add_run(f"\t{period}")


# --------------------------------------------------------------------------- #
# CV — DOCX
# --------------------------------------------------------------------------- #
def generate_cv_docx(data: dict, full_name: str) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.63)
        section.right_margin = Inches(0.63)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.1

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(full_name)
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_parts = _contact_line(data.get("contact"))
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(2)
        run = contact_p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9.5)
        _set_bottom_border(contact_p, color="000000")

    if data.get("summary"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(6)
        p.add_run(data["summary"])

    if data.get("experience"):
        _section_heading(doc, "Work Experiences")
        for exp in data["experience"]:
            _entry_header(doc, exp.get("company", ""), exp.get("period", ""))
            if exp.get("title"):
                role_p = doc.add_paragraph()
                role_p.paragraph_format.space_after = Pt(2)
                role_run = role_p.add_run(exp["title"])
                role_run.italic = True
            for bullet in exp.get("bullets", []):
                b = doc.add_paragraph(bullet, style="List Bullet")
                b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                b.paragraph_format.space_after = Pt(1)

    if data.get("education"):
        _section_heading(doc, "Education")
        for edu in data["education"]:
            org = edu.get("institution", "")
            _entry_header(doc, org, edu.get("year", ""))
            if edu.get("degree"):
                deg_p = doc.add_paragraph()
                deg_p.paragraph_format.space_after = Pt(2)
                deg_p.add_run(edu["degree"]).italic = True

    for key, heading in (
        ("certifications", "Certifications"),
        ("awards", "Awards"),
        ("projects", "Projects"),
    ):
        items = data.get(key)
        if items:
            _section_heading(doc, heading)
            for item in items:
                b = doc.add_paragraph(str(item), style="List Bullet")
                b.paragraph_format.space_after = Pt(1)

    if data.get("skills"):
        _section_heading(doc, "Skills")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(", ".join(data["skills"]))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_cover_letter_docx(
    content: str, full_name: str, contact: dict | None = None
) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(full_name)
    name_run.bold = True
    name_run.font.size = Pt(18)

    contact_parts = _contact_line(contact)
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(8)
        run = contact_p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9.5)
        _set_bottom_border(contact_p, color="000000")

    for paragraph in content.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            p = doc.add_paragraph(paragraph)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# CV — PDF (WeasyPrint)
# --------------------------------------------------------------------------- #
_CV_CSS = """
@page { size: A4; margin: 1.5cm 1.6cm; }
body { font-family: 'Times New Roman', Georgia, serif; font-size: 10.5pt; color: #000; line-height: 1.25; }
.name { text-align: center; font-size: 20pt; font-weight: bold; margin: 0 0 2px; }
.contact { text-align: center; font-size: 9.5pt; margin: 0 0 6px; padding-bottom: 6px; border-bottom: 1px solid #000; }
.contact a { color: #0563C1; text-decoration: none; }
.summary { text-align: justify; margin: 6px 0 4px; }
h2.section { font-size: 11.5pt; font-weight: bold; color: #1F4E79; text-transform: uppercase;
             letter-spacing: .3px; border-bottom: 1.2px solid #BFBFBF; padding-bottom: 2px; margin: 12px 0 5px; }
.entry-head { display: flex; justify-content: space-between; align-items: baseline; margin-top: 4px; }
.entry-org { font-weight: bold; }
.entry-period { white-space: nowrap; padding-left: 14px; }
.entry-role { font-style: italic; margin: 0 0 2px; }
ul { margin: 2px 0 4px; padding-left: 18px; }
li { margin-bottom: 2px; text-align: justify; }
.skills { text-align: justify; margin-top: 2px; }
"""


def _entry_head_html(org: str, period: str) -> str:
    return (
        "<div class='entry-head'>"
        f"<span class='entry-org'>{escape(org)}</span>"
        f"<span class='entry-period'>{escape(period)}</span>"
        "</div>"
    )


def _contact_html(contact: dict | None) -> str:
    parts = _contact_line(contact)
    if not parts:
        return ""
    rendered = []
    for part in parts:
        if part.startswith(("http://", "https://", "www.")):
            href = part if part.startswith("http") else f"https://{part}"
            rendered.append(f"<a href='{escape(href)}'>{escape(part)}</a>")
        else:
            rendered.append(escape(part))
    return f"<p class='contact'>{' | '.join(rendered)}</p>"


def generate_cv_pdf(data: dict, full_name: str) -> bytes | None:
    try:
        from weasyprint import HTML

        parts = [
            "<html><head><meta charset='utf-8'><style>",
            _CV_CSS,
            "</style></head><body>",
            f"<p class='name'>{escape(full_name)}</p>",
            _contact_html(data.get("contact")),
        ]

        if data.get("summary"):
            parts.append(f"<p class='summary'>{escape(data['summary'])}</p>")

        if data.get("experience"):
            parts.append("<h2 class='section'>Work Experiences</h2>")
            for exp in data["experience"]:
                parts.append(_entry_head_html(exp.get("company", ""), exp.get("period", "")))
                if exp.get("title"):
                    parts.append(f"<p class='entry-role'>{escape(exp['title'])}</p>")
                if exp.get("bullets"):
                    parts.append("<ul>")
                    parts.extend(f"<li>{escape(b)}</li>" for b in exp["bullets"])
                    parts.append("</ul>")

        if data.get("education"):
            parts.append("<h2 class='section'>Education</h2>")
            for edu in data["education"]:
                parts.append(_entry_head_html(edu.get("institution", ""), edu.get("year", "")))
                if edu.get("degree"):
                    parts.append(f"<p class='entry-role'>{escape(edu['degree'])}</p>")

        for key, heading in (
            ("certifications", "Certifications"),
            ("awards", "Awards"),
            ("projects", "Projects"),
        ):
            items = data.get(key)
            if items:
                parts.append(f"<h2 class='section'>{heading}</h2><ul>")
                parts.extend(f"<li>{escape(str(item))}</li>" for item in items)
                parts.append("</ul>")

        if data.get("skills"):
            parts.append("<h2 class='section'>Skills</h2>")
            parts.append(f"<p class='skills'>{escape(', '.join(data['skills']))}</p>")

        parts.append("</body></html>")
        return HTML(string="\n".join(parts)).write_pdf()
    except Exception:
        logger.exception("Failed to generate PDF")
        return None


def generate_cover_letter_pdf(
    content: str, full_name: str, contact: dict | None = None
) -> bytes | None:
    try:
        from weasyprint import HTML

        paragraphs = "".join(
            f"<p>{escape(p.strip())}</p>" for p in content.split("\n\n") if p.strip()
        )
        css = (
            "@page { size: A4; margin: 2cm 2.2cm; }"
            "body { font-family: 'Times New Roman', Georgia, serif; font-size: 11pt; color: #000; line-height: 1.5; }"
            ".name { text-align: center; font-size: 18pt; font-weight: bold; margin: 0 0 2px; }"
            ".contact { text-align: center; font-size: 9.5pt; margin: 0 0 14px; padding-bottom: 6px; border-bottom: 1px solid #000; }"
            "p { margin-bottom: 10px; text-align: justify; }"
        )
        html_content = (
            "<html><head><meta charset='utf-8'><style>"
            f"{css}</style></head><body>"
            f"<p class='name'>{escape(full_name)}</p>"
            f"{_contact_html(contact)}"
            f"{paragraphs}"
            "</body></html>"
        )
        return HTML(string=html_content).write_pdf()
    except Exception:
        logger.exception("Failed to generate cover letter PDF")
        return None
